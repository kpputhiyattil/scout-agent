"""Export ratings as CSV and a coach-facing Word summary.

The document leads with what the footage could and could not measure, because a
score computed from four seconds of fragmented tracking looks identical to one
computed from a full match unless you say so.

CLI:
  python -m scout.report.export --match <match_id> --out reports/
  python -m scout.report.export --latest
"""
from __future__ import annotations

import json
from datetime import date
from pathlib import Path

import pandas as pd

from scout.config import get_settings
from scout.db import Match, Player, Rating, get_session

ROLE_LABEL = {"GK": "Goalkeeper", "DEF": "Defender", "MID": "Midfielder", "ATT": "Attacker"}
ROLE_ORDER = ["GK", "DEF", "MID", "ATT"]

# A readable document, not a data dump — the CSV carries the full detail.
MAX_ROWS_PER_BLOCK = 8      # per team-and-position table
MAX_NOTES = 15


def _team_names(df: pd.DataFrame) -> list[str]:
    """Identified teams first, an unidentified group ('?') last if present."""
    teams = [t for t in df.team.dropna().unique() if str(t) not in ("?", "")]
    unknown = [t for t in df.team.dropna().unique() if str(t) in ("?", "")]
    return sorted(teams) + sorted(unknown)


def _quality(match_id: str) -> dict:
    p = get_settings().match_dir(match_id) / "quality.json"
    return json.loads(p.read_text()) if p.exists() else {}


def _mode(match_id: str) -> str:
    p = get_settings().match_dir(match_id) / "projection.json"
    return json.loads(p.read_text()).get("mode", "pitch") if p.exists() else "pitch"


def load_match(match_id: str) -> tuple[Match, list[Player], dict[int, Rating]]:
    with get_session() as db:
        match = db.get(Match, match_id)
        if match is None:
            raise ValueError(f"No such match: {match_id}")
        players = db.query(Player).filter_by(match_id=match_id).all()
        ratings = {r.player_id: r for r in db.query(Rating).filter_by(match_id=match_id).all()}
        db.expunge_all()
    return match, players, ratings


def latest_match_id() -> str:
    with get_session() as db:
        m = db.query(Match).order_by(Match.created_at.desc()).first()
        if m is None:
            raise ValueError("No matches in the database yet")
        return m.id


def ratings_frame(match_id: str) -> pd.DataFrame:
    """One row per player: identity, position, rating, sub-scores and raw evidence."""
    _, players, ratings = load_match(match_id)
    rows = []
    for p in players:
        r = ratings.get(p.id)
        row = {
            "player": p.name,
            "team": p.team,
            "jersey": p.jersey,
            "position": ROLE_LABEL.get(p.role, p.role),
            "position_confidence": round(p.role_confidence or 0, 2),
            "minutes": round(p.minutes, 1),
            "rating": r.overall if r else None,
        }
        if r:
            row.update({f"score_{k}": v for k, v in (r.sub_scores or {}).items()})
            row.update({f"kpi_{k}": v for k, v in (r.evidence or {}).items()})
            row["note"] = r.note
        rows.append(row)
    df = pd.DataFrame(rows)
    return df.sort_values("rating", ascending=False, na_position="last").reset_index(drop=True)


def report_frame(match_id: str) -> pd.DataFrame:
    """Presentation view: grouped by team, then position (GK→DEF→MID→ATT),
    best player first inside each block, with ranks precomputed.

    Sorting this way means the file answers the questions a coach actually asks —
    "who were my best midfielders?" — by reading top-down, with no spreadsheet work.
    """
    df = ratings_frame(match_id)
    if df.empty:
        return df

    role_key = {ROLE_LABEL[r]: i for i, r in enumerate(ROLE_ORDER)}
    df["_role_order"] = df.position.map(role_key).fillna(len(ROLE_ORDER))
    # unidentified teams and unrated players belong at the bottom, not the top
    df["_team_order"] = df.team.map(lambda t: (t in (None, "?", ""), str(t)))
    df["_unrated"] = df.rating.isna()
    df = df.sort_values(["_unrated", "_team_order", "_role_order", "rating"],
                        ascending=[True, True, True, False], na_position="last")

    df["rank_in_position"] = df.groupby(["team", "position"]).cumcount() + 1
    df["rank_in_team"] = (df.groupby("team")["rating"]
                            .rank(ascending=False, method="min").astype("Int64"))
    df["rank_overall"] = df["rating"].rank(ascending=False, method="min").astype("Int64")

    lead = ["team", "position", "rank_in_position", "player", "jersey", "rating",
            "minutes", "rank_in_team", "rank_overall", "position_confidence"]
    scores = sorted(c for c in df.columns if c.startswith("score_"))
    kpis = sorted(c for c in df.columns if c.startswith("kpi_"))
    tail = [c for c in ("note",) if c in df.columns]
    ordered = [c for c in lead + scores + kpis + tail if c in df.columns]

    out = df[ordered].reset_index(drop=True)

    def header(c: str) -> str:
        """Title Case reads better in Excel than snake_case; keep 'KPI' upper."""
        if c.startswith("kpi_"):
            return "KPI: " + c[4:].replace("_", " ").title()
        if c.startswith("score_"):
            return "Score: " + c[6:].replace("_", " ").title()
        return c.replace("_", " ").title()

    return out.rename(columns=header)


def export_csv(match_id: str, out_path: str | Path) -> Path:
    """Coach-readable CSV: grouped by team and position, best first."""
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    report_frame(match_id).to_csv(out, index=False)
    return out


def best_by_position(match_id: str) -> dict[str, tuple[Player, Rating] | None]:
    _, players, ratings = load_match(match_id)
    best = {}
    for role in ROLE_ORDER:
        cands = [(p, ratings[p.id]) for p in players
                 if p.role == role and p.id in ratings and ratings[p.id].overall is not None]
        best[role] = max(cands, key=lambda x: x[1].overall) if cands else None
    return best


def _caveats(match_id: str) -> list[str]:
    """Plain-language limits of this analysis — the part a coach must read."""
    q, mode, out = _quality(match_id), _mode(match_id), []
    if mode == "relative":
        out.append(
            "No fixed pitch reference was available (moving camera or partial pitch view), "
            "so ratings use on-ball actions only: touches, passing, duels, possession won "
            "and lost, and shots. Distance covered, speed and positioning were NOT measured "
            "and are excluded from every score.")
    if q:
        raw, rated = q.get("n_raw_tracks", 0), max(q.get("n_rated", 0), 1)
        if raw > 4 * rated:
            out.append(
                f"Tracking was heavily fragmented: {raw} separate identities were created for "
                f"about {rated} players. Camera cuts, panning and players blocking one another "
                "make the tracker lose people, so a single child can appear as many short "
                "tracks. Treat individual scores as indicative only.")
        if q.get("n_unidentified_dropped"):
            out.append(
                f"{q['n_unidentified_dropped']} tracked identities had no readable jersey "
                "number and were excluded: a rating no one can attribute to a specific child "
                "is not useful. Clearer footage of numbers, or a roster CSV, increases coverage.")
        if q.get("n_over_squad_size_dropped"):
            out.append(
                f"{q['n_over_squad_size_dropped']} identities beyond a plausible squad size "
                "were excluded — these are usually the same child counted twice under "
                "different numbers after a misread.")
    med = q.get("median_minutes", 0)
    if med and med < 15:
        out.append(
            f"Median observed time per player is {med:.1f} minutes. Per-90 figures extrapolated "
            "from short samples are noisy — compare players with similar minutes.")
    out.append(
        "Ratings are percentiles within this squad, not absolute ability. A 70 means "
        "'better than most peers in this match on these measures', not '70% of a professional'.")
    return out


def export_docx(match_id: str, out_path: str | Path) -> Path:
    """Coach-facing Word summary: caveats, squad table, standouts, per-player notes."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError as e:  # pragma: no cover - depends on install extras
        raise RuntimeError(
            "python-docx is required for Word export. Install it with "
            '`pip install python-docx` (or `pip install -e ".[ui]"`).') from e

    match, players, ratings = load_match(match_id)
    q, mode = _quality(match_id), _mode(match_id)
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("Match scouting report", level=0)
    sub = doc.add_paragraph()
    sub.add_run(f"Generated {date.today():%d %B %Y}").italic = True
    doc.add_paragraph(f"Source: {match.source}")
    doc.add_paragraph(
        f"Duration: {match.duration_s / 60:.1f} min   |   "
        f"Measurement mode: {'camera-relative (on-ball only)' if mode == 'relative' else 'full pitch metrics'}   |   "
        f"Players rated: {q.get('n_rated', len(players))}")

    doc.add_heading("How to read this report", level=1)
    for c in _caveats(match_id):
        doc.add_paragraph(c, style="List Bullet")

    doc.add_heading("Standouts", level=1)
    p = doc.add_paragraph()
    p.add_run("Best overall by position (both teams)").bold = True
    for role, best in best_by_position(match_id).items():
        line = doc.add_paragraph(style="List Bullet")
        line.add_run(f"{ROLE_LABEL[role]}: ").bold = True
        if best:
            player, rating = best
            line.add_run(f"{player.name} (Team {player.team}) — {rating.overall:.0f}/100, "
                         f"{player.minutes:.1f} min observed")
        else:
            line.add_run("no player identified in this position")

    df_all = ratings_frame(match_id)
    for team in _team_names(df_all):
        block = df_all[df_all.team == team].sort_values("rating", ascending=False,
                                                        na_position="last").head(3)
        if block.empty:
            continue
        p = doc.add_paragraph()
        p.add_run(f"Top rated — Team {team}").bold = True
        for _, row in block.iterrows():
            line = doc.add_paragraph(style="List Bullet")
            line.add_run(f"{row.player} ({row.position}) — "
                         f"{'—' if pd.isna(row.rating) else f'{row.rating:.0f}'}/100")

    doc.add_heading("Ratings by team and position", level=1)
    doc.add_paragraph(
        "Within each position the best-rated player is listed first. The accompanying "
        "CSV holds every player and every raw measurement.")

    df = ratings_frame(match_id)
    score_cols = [c for c in df.columns if c.startswith("score_")]
    headers = ["player", "jersey", "minutes", "rating"] + score_cols
    widths = [Inches(1.6), Inches(0.6), Inches(0.7), Inches(0.7)]
    widths += [Inches(0.85)] * len(score_cols)

    def _table(block: pd.DataFrame) -> None:
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        table.autofit = False
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.width = widths[i]
            cell.text = h.replace("score_", "").replace("_", " ").title()
            for r in cell.paragraphs[0].runs:
                r.bold = True
                r.font.size = Pt(8)
        for _, row in block.iterrows():
            cells = table.add_row().cells
            for i, h in enumerate(headers):
                v = row.get(h)
                cells[i].width = widths[i]
                cells[i].text = ("—" if pd.isna(v)
                                 else f"{v:.1f}" if isinstance(v, float) else str(v))
                for r in cells[i].paragraphs[0].runs:
                    r.font.size = Pt(8)

    for team in _team_names(df):
        doc.add_heading(f"Team {team}", level=2)
        team_df = df[df.team == team]
        for role in ROLE_ORDER:
            block = (team_df[team_df.position == ROLE_LABEL[role]]
                     .sort_values("rating", ascending=False, na_position="last"))
            if block.empty:
                continue
            p = doc.add_paragraph()
            p.add_run(f"{ROLE_LABEL[role]}s ({len(block)})").bold = True
            shown = block.head(MAX_ROWS_PER_BLOCK)
            _table(shown)
            if len(block) > len(shown):
                more = doc.add_paragraph()
                mr = more.add_run(f"…and {len(block) - len(shown)} more in the CSV")
                mr.italic = True
                mr.font.size = Pt(8)

    doc.add_heading("Player notes", level=1)
    id_of = {p.id: p for p in players}
    ranked = sorted((r for r in ratings.values() if r.overall is not None),
                    key=lambda r: r.overall, reverse=True)
    if len(ranked) > MAX_NOTES:
        doc.add_paragraph(f"Notes for the top {MAX_NOTES} rated players.")
        ranked = ranked[:MAX_NOTES]
    for r in ranked:
        p = id_of.get(r.player_id)
        if p is None:
            continue
        h = doc.add_paragraph()
        h.add_run(f"{p.name} — {ROLE_LABEL.get(p.role, p.role)} — {r.overall:.0f}/100").bold = True
        if r.note:
            doc.add_paragraph(r.note)
        if r.evidence:
            ev = ", ".join(f"{k.replace('_', ' ')}: {v}" for k, v in list(r.evidence.items())[:8])
            small = doc.add_paragraph()
            run = small.add_run(f"Measured: {ev}")
            run.italic = True
            run.font.size = Pt(9)

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run("Generated by ScoutTrainer — ratings are percentile-based within this squad "
                      "and intended to support, not replace, a coach's judgement.")
    fr.italic = True
    fr.font.size = Pt(8)

    doc.save(out)
    return out


def export_all(match_id: str | None = None, out_dir: str | Path | None = None) -> dict[str, Path]:
    """Write both CSV and Word summary; returns {'csv': path, 'docx': path}."""
    match_id = match_id or latest_match_id()
    out_dir = Path(out_dir) if out_dir else get_settings().match_dir(match_id) / "reports"
    out_dir.mkdir(parents=True, exist_ok=True)
    paths = {"csv": export_csv(match_id, out_dir / f"ratings_{match_id}.csv")}
    try:
        paths["docx"] = export_docx(match_id, out_dir / f"scouting_report_{match_id}.docx")
    except RuntimeError as e:
        print(f"Word export skipped: {e}")
    return paths


def main():
    import argparse
    ap = argparse.ArgumentParser(description="Export ScoutTrainer ratings")
    g = ap.add_mutually_exclusive_group(required=True)
    g.add_argument("--match", help="Match id")
    g.add_argument("--latest", action="store_true", help="Most recently added match")
    ap.add_argument("--out", help="Output directory (default: the match's reports/ folder)")
    a = ap.parse_args()
    for kind, path in export_all(None if a.latest else a.match, a.out).items():
        print(f"{kind}: {path}")


if __name__ == "__main__":
    main()
